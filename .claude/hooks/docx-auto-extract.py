"""Hook: auto-extract .docx files when Read tool fails on them.

PostToolUseFailure hook for Read — when Read fails on a .docx file,
extract text via python-docx and save to cache dir, then print the
path so Claude knows to re-read the extracted .txt version.
"""
import json
import os
import sys

try:
    data = json.load(sys.stdin)
    path = data.get("tool_input", {}).get("file_path", "")

    if not path.lower().endswith((".docx", ".doc")):
        sys.exit(0)

    if not os.path.isfile(path):
        sys.exit(0)

    # Only handle .docx (python-docx doesn't support old .doc)
    if path.lower().endswith(".doc") and not path.lower().endswith(".docx"):
        print(f"[hook] {os.path.basename(path)} is old .doc format — python-docx can't read it. Open in Word or convert to .docx first.")
        sys.exit(0)

    from docx import Document

    cache_dir = os.path.join(os.path.expanduser("~"), "VoxCore", ".claude", "cache", "docx_extractions")
    os.makedirs(cache_dir, exist_ok=True)

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)

    # Also extract tables if present
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + "\t".join(cell.text for cell in row.cells)

    basename = os.path.splitext(os.path.basename(path))[0]
    out_path = os.path.join(cache_dir, f"{basename}.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"[hook] Auto-extracted .docx to readable text: {out_path}")
    print(f"[hook] Use Read on that .txt path to see the contents.")

except Exception as e:
    # Don't block on errors — just note it
    print(f"[hook] docx extraction failed: {e}", file=sys.stderr)
