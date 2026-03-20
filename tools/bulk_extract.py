#!/usr/bin/env python3
"""Batch-extract text from PDFs, .docx, .eml, and .msg files in a directory tree.

Produces a single searchable text dump or per-file outputs. Agents read the
dump instead of opening files one-by-one.

Usage:
    python tools/bulk_extract.py <directory>
    python tools/bulk_extract.py <directory> --output extracted.txt
    python tools/bulk_extract.py <directory> --format json --output extracted.json
    python tools/bulk_extract.py <directory> --per-file --outdir ./extracted/
    python tools/bulk_extract.py <directory> --types .pdf,.docx

Options:
    --output    Single output file (default: stdout)
    --format    Output format: text or json (default: text)
    --per-file  Write one .txt per source file into --outdir
    --outdir    Directory for per-file output (default: ./extracted/)
    --types     Comma-separated extensions to process (default: .pdf,.docx,.eml,.msg,.doc)
    --workers   Parallel workers (default: 4)

Requirements:
    pip install python-docx PyPDF2
    (eml/msg parsing uses stdlib email module — no extra deps)
"""
import argparse
import email
import json
import os
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from email import policy


def extract_pdf(path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"--- Page {i+1} ---\n{text}")
        return "\n\n".join(pages) if pages else "(no extractable text)"
    except ImportError:
        return "(PyPDF2 not installed — pip install PyPDF2)"
    except Exception as e:
        return f"(PDF extraction error: {e})"


def extract_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts) if parts else "(empty document)"
    except ImportError:
        return "(python-docx not installed — pip install python-docx)"
    except Exception as e:
        return f"(DOCX extraction error: {e})"


def extract_eml(path):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            msg = email.message_from_file(f, policy=policy.default)

        parts = []
        parts.append(f"From: {msg.get('From', '(unknown)')}")
        parts.append(f"To: {msg.get('To', '(unknown)')}")
        parts.append(f"Date: {msg.get('Date', '(unknown)')}")
        parts.append(f"Subject: {msg.get('Subject', '(no subject)')}")
        parts.append("")

        body = msg.get_body(preferencelist=("plain", "html"))
        if body:
            content = body.get_content()
            if isinstance(content, bytes):
                content = content.decode("utf-8", errors="replace")
            parts.append(content)
        else:
            # Walk parts for multipart messages
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        parts.append(payload.decode("utf-8", errors="replace"))

        # List attachments
        attachments = []
        for part in msg.iter_attachments():
            fn = part.get_filename()
            if fn:
                attachments.append(fn)
        if attachments:
            parts.append(f"\nAttachments: {', '.join(attachments)}")

        return "\n".join(parts) if parts else "(empty email)"
    except Exception as e:
        return f"(EML extraction error: {e})"


def extract_msg(path):
    # .msg is Outlook format — try extract_msg library, fall back to binary notice
    try:
        import extract_msg
        msg = extract_msg.Message(path)
        parts = [
            f"From: {msg.sender}",
            f"To: {msg.to}",
            f"Date: {msg.date}",
            f"Subject: {msg.subject}",
            "",
            msg.body or "(no body)",
        ]
        msg.close()
        return "\n".join(parts)
    except ImportError:
        return "(extract-msg not installed — pip install extract-msg)"
    except Exception as e:
        return f"(MSG extraction error: {e})"


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_docx,  # python-docx sometimes handles .doc
    ".eml": extract_eml,
    ".msg": extract_msg,
}


def process_file(filepath, root):
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()
    extractor = EXTRACTORS.get(ext)
    if not extractor:
        return None

    relpath = os.path.relpath(filepath, root).replace("\\", "/")
    text = extractor(filepath)
    size = os.path.getsize(filepath)

    return {
        "path": relpath,
        "ext": ext,
        "size": size,
        "text": text,
    }


def main():
    parser = argparse.ArgumentParser(description="Batch-extract text from PDFs, DOCX, EML, MSG")
    parser.add_argument("directory", help="Root directory to scan")
    parser.add_argument("--output", "-o", help="Output file (default: stdout)")
    parser.add_argument("--format", choices=["text", "json"], default="text", help="Output format")
    parser.add_argument("--per-file", action="store_true", help="Write one .txt per source file")
    parser.add_argument("--outdir", default="./extracted/", help="Dir for per-file output")
    parser.add_argument("--types", default=".pdf,.docx,.eml,.msg,.doc",
                        help="Extensions to process")
    parser.add_argument("--workers", type=int, default=4, help="Parallel workers")
    args = parser.parse_args()

    root = os.path.expanduser(args.directory)
    if not os.path.isdir(root):
        print(f"Error: {root} is not a directory", file=sys.stderr)
        sys.exit(1)

    include_types = set(args.types.split(","))

    # Collect files
    files = []
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            _, ext = os.path.splitext(fn)
            if ext.lower() in include_types:
                files.append(os.path.join(dirpath, fn))

    if not files:
        print(f"No matching files found in {root}", file=sys.stderr)
        sys.exit(0)

    print(f"Processing {len(files)} files with {args.workers} workers...", file=sys.stderr)

    # Process in parallel
    results = []
    failed = 0
    with ThreadPoolExecutor(max_workers=args.workers) as pool:
        futures = {pool.submit(process_file, f, root): f for f in files}
        for future in as_completed(futures):
            result = future.result()
            if result:
                results.append(result)
            else:
                failed += 1

    # Sort by path for deterministic output
    results.sort(key=lambda r: r["path"])
    print(f"Extracted {len(results)} files ({failed} skipped)", file=sys.stderr)

    # Per-file output mode
    if args.per_file:
        os.makedirs(args.outdir, exist_ok=True)
        for r in results:
            out_name = r["path"].replace("/", "__").rsplit(".", 1)[0] + ".txt"
            out_path = os.path.join(args.outdir, out_name)
            with open(out_path, "w", encoding="utf-8", newline="\n") as f:
                f.write(f"Source: {r['path']}\nSize: {r['size']:,} bytes\n\n{r['text']}")
        print(f"Wrote {len(results)} files to {args.outdir}", file=sys.stderr)
        return

    # Single output mode
    if args.format == "json":
        text = json.dumps(results, indent=2, ensure_ascii=False)
    else:
        parts = []
        for r in results:
            header = f"{'='*80}\nFILE: {r['path']} ({r['size']:,} bytes)\n{'='*80}"
            parts.append(f"{header}\n{r['text']}")
        text = "\n\n".join(parts)

    if args.output:
        with open(args.output, "w", encoding="utf-8", newline="\n") as f:
            f.write(text)
        print(f"Wrote {len(text):,} bytes to {args.output}", file=sys.stderr)
    else:
        print(text)


if __name__ == "__main__":
    main()
